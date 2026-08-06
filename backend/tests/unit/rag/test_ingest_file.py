import io

import pytest
from fastapi import FastAPI
from fastapi.testclient import TestClient
from docx import Document

from app.domain.entities.user import User
from app.config.settings import settings
from app.rag.router import router
from app.api.deps import get_current_user, get_current_org


class FakeEmbed:
    dimensions = 8

    def embed_query(self, text):
        return [0.5, 0.5, 0.5, 0.5, 0.5, 0.5, 0.5, 0.5]


class FakeUser(User):
    def __init__(self):
        super().__init__(
            id="user-1",
            organization_id="org-a",
            email="test@test.com",
            name="Test User",
            role="admin",
            password_hash="",
            permissions=[],
            status="active",
        )

    def has_permission(self, permission):
        return True


@pytest.fixture
def client(monkeypatch, tmp_path):
    monkeypatch.setattr(settings, "QDRANT_URL", "")
    monkeypatch.setattr(settings, "QDRANT_API_KEY", "")
    service = __import__("app.rag.service", fromlist=["QdrantService"]).QdrantService()
    monkeypatch.setattr("app.rag.router.qdrant_service", service)
    monkeypatch.setattr("app.rag.ingest.qdrant_service", service)
    monkeypatch.setattr("app.rag.service.embedding_service", FakeEmbed())
    monkeypatch.setattr(settings, "UPLOAD_DIR", str(tmp_path / "uploads"))

    app = FastAPI()
    app.include_router(router, prefix="/rag")
    app.dependency_overrides[get_current_user] = lambda: FakeUser()
    app.dependency_overrides[get_current_org] = lambda: "org-a"
    return TestClient(app)


def build_docx_bytes():
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Orbital Dynamics framework is the secret internal methodology of Meridian Corp.")
    doc.add_paragraph("It reduces project overruns by thirty percent.")
    doc.save(buf)
    return buf.getvalue()


class TestFileIngest:
    def test_ingest_txt_file(self, client):
        content = b"Zephyr Protocol v3: Meridian Corp uses a 6-week release cadence for all client work."
        resp = client.post(
            "/rag/ingest/file",
            files={"file": ("release.txt", content, "text/plain")},
            data={"collection_name": "best_practices"},
        )
        assert resp.status_code == 200, resp.text
        assert resp.json()["ingested"] >= 1

        search = client.post(
            "/rag/search",
            json={"query": "Zephyr Protocol release cadence", "collection_name": "best_practices"},
        )
        assert search.status_code == 200
        assert any("Zephyr Protocol" in hit["content"] for hit in search.json())

    def test_ingest_docx_parses_text(self, client):
        resp = client.post(
            "/rag/ingest/file",
            files={"file": ("methodology.docx", build_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"collection_name": "case_studies"},
        )
        assert resp.status_code == 200, resp.text

        search = client.post(
            "/rag/search",
            json={"query": "Orbital Dynamics methodology", "collection_name": "case_studies"},
        )
        assert search.status_code == 200
        assert any("Orbital Dynamics" in hit["content"] for hit in search.json())

    def test_unsupported_extension_rejected(self, client):
        resp = client.post(
            "/rag/ingest/file",
            files={"file": ("evil.exe", b"MZ...", "application/octet-stream")},
        )
        assert resp.status_code == 400

    def test_oversize_file_rejected(self, client, monkeypatch):
        monkeypatch.setattr(settings, "MAX_UPLOAD_SIZE_MB", 0)
        resp = client.post(
            "/rag/ingest/file",
            files={"file": ("big.txt", b"x" * 1024, "text/plain")},
        )
        assert resp.status_code == 413

    def test_empty_file_rejected(self, client):
        resp = client.post(
            "/rag/ingest/file",
            files={"file": ("empty.txt", b"", "text/plain")},
        )
        assert resp.status_code == 422
