import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.export.renderers.base import BaseRenderer
from app.export.renderers.common import iter_renderable_sections, build_section_blocks

logger = logging.getLogger("proposalcraft.export.docx")


class DOCXRenderer(BaseRenderer):
    extension = ".docx"

    def render(self, proposal: dict, output_path: str) -> str:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

        self._add_cover(doc, proposal.get("metadata", {}))

        first = True
        for key, title, data in iter_renderable_sections(proposal):
            if not first:
                doc.add_page_break()
            first = False
            self._add_section(doc, title, build_section_blocks(key, data))

        path = Path(output_path).with_suffix(".docx")
        doc.save(str(path))
        logger.info("DOCX exported to %s", path)
        return str(path)

    def _add_cover(self, doc: Document, meta: dict):
        for _ in range(6):
            doc.add_paragraph("")

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(meta.get("proposal_title", "Proposal"))
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(26, 86, 219)
        run.bold = True

        subtitle = meta.get("subtitle")
        if subtitle:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(subtitle)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph("")
        for label, key in [("Prepared for", "prepared_for"), ("Prepared by", "prepared_by"),
                           ("Date", "date"), ("Version", "version")]:
            val = meta.get(key)
            if val:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"{label}: {val}")
                run.font.size = Pt(11)

    def _add_section(self, doc: Document, title: str, blocks: list[tuple]):
        heading = doc.add_heading(title, level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(26, 86, 219)

        for block in blocks:
            kind = block[0]
            if kind == "h3":
                h = doc.add_heading(block[1], level=2)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(26, 86, 219)
            elif kind == "p":
                doc.add_paragraph(block[1])
            elif kind == "bullets":
                for item in block[1]:
                    doc.add_paragraph(item, style="List Bullet")
            elif kind == "table":
                headers, rows = block[1], block[2]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Light Grid Accent 1"
                for i, htext in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = htext
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                for row in rows:
                    cells = table.add_row().cells
                    for i, value in enumerate(row):
                        if i < len(cells):
                            cells[i].text = str(value)
